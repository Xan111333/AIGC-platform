import os
import io
import zipfile
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def text_to_pdf(text: str, title: str = "Generated Text") -> bytes:
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    
    c.setFont("Helvetica-Bold", 16)
    c.drawCentredString(width / 2, height - 50, title)
    
    c.setFont("Helvetica", 12)
    text_lines = text.split('\n')
    y = height - 100
    line_height = 20
    
    for line in text_lines:
        if y < 50:
            c.showPage()
            c.setFont("Helvetica", 12)
            y = height - 50
        
        words = line.split(' ')
        current_line = ""
        for word in words:
            test_line = current_line + " " + word if current_line else word
            if c.stringWidth(test_line, "Helvetica", 12) < width - 100:
                current_line = test_line
            else:
                c.drawString(50, y, current_line)
                y -= line_height
                current_line = word
        
        if current_line:
            c.drawString(50, y, current_line)
            y -= line_height
    
    c.showPage()
    c.save()
    
    buffer.seek(0)
    return buffer.read()

def text_to_word(text: str, title: str = "Generated Text") -> bytes:
    doc = Document()
    
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph(text)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

def images_to_zip(image_urls: list) -> bytes:
    buffer = io.BytesIO()
    
    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for i, url in enumerate(image_urls):
            try:
                import requests
                response = requests.get(url, timeout=10)
                response.raise_for_status()
                
                filename = f"image_{i+1}.png"
                zipf.writestr(filename, response.content)
            except Exception as e:
                print(f"Failed to download {url}: {e}")
    
    buffer.seek(0)
    return buffer.read()

def generate_filename(prefix: str, extension: str) -> str:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{prefix}_{timestamp}.{extension}"